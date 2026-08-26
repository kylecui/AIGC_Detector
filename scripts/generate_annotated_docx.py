"""Generate annotated Word doc: pandoc for clean format, then python-docx for highlights."""
import re
import subprocess
from pathlib import Path
from docx import Document
from docx.enum.text import WD_COLOR_INDEX

MD_PATH = "docs/patent-application/专利技术交底书.md"
DOCX_PATH = "outputs/patent-pdf/专利技术交底书_批注回应版.docx"
WORK_DIR = Path(DOCX_PATH).parent
WORK_DIR.mkdir(parents=True, exist_ok=True)

# Step 1: pandoc markdown to docx (clean formatting)
content = Path(MD_PATH).read_text(encoding="utf-8")
content = content.replace("../diagrams/", str(Path("docs/diagrams/").resolve()).replace("\\", "/") + "/")
tmp_md = WORK_DIR / "tmp_for_pandoc.md"
tmp_md.write_text(content, encoding="utf-8")

print("Step 1: pandoc converting...")
result = subprocess.run(
    ["pandoc", str(tmp_md), "-o", str(DOCX_PATH),
     "--from", "markdown", "--to", "docx",
     "--resource-path", str(Path("docs/diagrams/").resolve())],
    capture_output=True, text=True, timeout=60
)
if result.returncode != 0:
    print(f"pandoc error: {result.stderr}")
    exit(1)
print(f"pandoc OK")

# Step 2: python-docx post-processing - add yellow highlights
HIGHLIGHT_KEYWORDS = [
    ("当语言识别置信度低于0.80", "#0"),
    ("Burstiness", "#1"),
    ("输入文本预处理", "#2,#9"),
    ("文本清洗与格式归一化", "#2,#9"),
    ("中英文分句规则", "#2,#9"),
    ("分词工具选型", "#2,#9"),
    ("边界场景处理", "#2,#9"),
    ("融合置信度", "#3"),
    ("中文仲裁块", "#4"),
    ("反向场景", "#4"),
    ("缓存数据对齐规则", "#5"),
    ("缓存生命周期与降级处理", "#5"),
    ("句首句法重复率", "#6,#10"),
    ("段落模板得分", "#6,#10"),
    ("全量数据打乱后随机采样", "#7"),
    ("不超过2个百分点", "#8,#11"),
    ("Acc_new", "#8,#11"),
    ("M7模糊语词表", "补充"),
    ("M8话语模板词表", "补充"),
    ("Sigmoid", "补充"),
    ("0.854", "补充"),
]

print("Step 2: adding highlights...")
doc = Document(DOCX_PATH)

count = 0
for keyword, tag in HIGHLIGHT_KEYWORDS:
    for paragraph in doc.paragraphs:
        if keyword in paragraph.text:
            for run in paragraph.runs:
                if keyword in run.text:
                    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                    count += 1
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if keyword in paragraph.text:
                        for run in paragraph.runs:
                            if keyword in run.text:
                                run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                                count += 1

print(f"Highlighted {count} runs")

# Step 3: Add summary page
doc.add_page_break()
doc.add_heading("批注回应高亮说明", level=1)

p = doc.add_paragraph()
r1 = p.add_run("本文档中所有回应专利代理师批注意见的内容均以")
r2 = p.add_run("黄色高亮")
r2.font.highlight_color = WD_COLOR_INDEX.YELLOW
r2.bold = True
p.add_run("标注。请在Word中搜索黄色高亮段落查看所有修改内容。")

doc.add_heading("批注对照表", level=2)

summary_data = [
    ("#0", "语言路由器降级策略缺失", "5.1", "置信度<0.80或中英混合时按字符统计判定主语言"),
    ("#1", "突发性术语混淆", "3.1", "Burstiness公式LaTeX化"),
    ("#2,#9", "预处理模块缺失", "5.0", "新增完整预处理(清洗/分句/分词/边界场景)"),
    ("#3", "提前退出置信度不清", "5.2", "C=w_stat*|p-0.5|+w_ling*|p-0.5|, C>0.49"),
    ("#4", "仲裁规则不完整", "5.2", "如实描述正向仲裁+反向由加权集成处理"),
    ("#5", "缓存对齐/生命周期缺失", "5.3", "句子索引对齐+请求级缓存+NaN降级"),
    ("#6,#10", "M4/S2公式不充分", "5.4", "M4 Jaccard公式+S2模板得分公式"),
    ("#7", "训练采样策略缺失", "5.6", "全量打乱随机采样"),
    ("#8,#11", "无退化缺乏量化", "5.6", "Acc_new>=Acc_old-2%"),
    ("补充", "Binoculars映射缺失", "5.7", "p_ai=1/(1+exp(5(score/s0-1)))"),
    ("补充", "M7/M8词表不完整", "5.4", "完整词表:M7(36词) M8(23词)"),
    ("补充", "权利要求脱节", "9", "claim1加预处理; claim5量化<=2%"),
]

table = doc.add_table(rows=1, cols=4)
try:
    table.style = "Table Grid"
except Exception:
    pass  # pandoc-generated docx may not have Table Grid style
hdr = table.rows[0].cells
hdr[0].text = "批注#"
hdr[1].text = "问题"
hdr[2].text = "位置"
hdr[3].text = "回应内容"

for row_data in summary_data:
    row = table.add_row()
    for i, val in enumerate(row_data):
        row.cells[i].text = val

doc.save(DOCX_PATH)
size_kb = Path(DOCX_PATH).stat().st_size / 1024
print(f"\nDone: {DOCX_PATH} ({size_kb:.0f} KB)")
